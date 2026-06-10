"""
Multi-format Document Generators
Creates professional documents in PDF, Word (.docx), and Google Docs formats.
"""
from __future__ import annotations

import os
from datetime import datetime
from typing import Dict

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from src.services.render_helpers import (
    is_bullet_block,
    is_markdown_table,
    normalize_unicode_text,
    parse_markdown_table,
    placeholder_items,
)


class WordDocumentGenerator:
    """Generates professional MS Word (.docx) monographs with proper formatting."""

    def __init__(self, output_dir: str = "data/monographs"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def generate_word_monograph(self, monograph_data: Dict, output_filename: str = None) -> str:
        if not output_filename:
            molecule_name = monograph_data.get("molecule_name", "Monograph").replace(" ", "_")
            output_filename = f"{molecule_name}_monograph_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

        filepath = os.path.join(self.output_dir, output_filename)
        doc = Document()

        for section in doc.sections:
            section.top_margin = Inches(0.85)
            section.bottom_margin = Inches(0.85)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        self._add_title_page(doc, monograph_data)
        doc.add_page_break()
        self._add_table_of_contents(doc, monograph_data)
        doc.add_page_break()

        if monograph_data.get("executive_summary"):
            self._add_section(doc, "Executive Summary", monograph_data["executive_summary"])
            doc.add_page_break()

        for section_name, section_content in monograph_data.get("sections", {}).items():
            self._add_section(doc, section_name.replace("_", " ").title(), section_content)

        if monograph_data.get("draft_placeholders"):
            doc.add_page_break()
            self._add_placeholders(doc, monograph_data["draft_placeholders"])

        if monograph_data.get("references"):
            doc.add_page_break()
            self._add_references(doc, monograph_data["references"])

        doc.add_page_break()
        self._add_disclaimer(doc)

        doc.save(filepath)
        print(f"[OK] Word document generated: {filepath}")
        return filepath

    def _add_title_page(self, doc: Document, monograph_data: Dict) -> None:
        for _ in range(4):
            doc.add_paragraph()

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(normalize_unicode_text(monograph_data.get("molecule_name", "Product Monograph")))
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(26, 26, 26)

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run("Product Monograph")
        subtitle_run.font.size = Pt(14)
        subtitle_run.font.color.rgb = RGBColor(100, 100, 100)

        doc.add_paragraph()
        metadata = [
            f"Generated: {datetime.now().strftime('%B %d, %Y')}",
            "Status: Draft - Auto-Generated",
            f"Compliance Score: {monograph_data.get('validation', {}).get('overall_compliance_score', 'N/A')}%",
            f"Total Sections: {len(monograph_data.get('sections', {}))}",
        ]
        for line in metadata:
            p = doc.add_paragraph(normalize_unicode_text(line))
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)

        notice = doc.add_paragraph()
        notice.alignment = WD_ALIGN_PARAGRAPH.LEFT
        notice_run = notice.add_run(
            "IMPORTANT NOTICE: This document is a draft generated for medical and regulatory review only. "
            "It must be checked against source documents before any clinical or publication use."
        )
        notice_run.font.size = Pt(10)
        notice_run.font.italic = True

    def _add_table_of_contents(self, doc: Document, monograph_data: Dict) -> None:
        heading = doc.add_heading("Table of Contents", level=1)
        heading.runs[0].font.color.rgb = RGBColor(44, 90, 160)
        for i, section_name in enumerate(monograph_data.get("sections", {}).keys(), 1):
            p = doc.add_paragraph(f"{i}. {section_name.replace('_', ' ').title()}")
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_after = Pt(2)

    def _add_section(self, doc: Document, section_title: str, content: str) -> None:
        heading = doc.add_heading(normalize_unicode_text(section_title), level=1)
        heading.runs[0].font.color.rgb = RGBColor(44, 90, 160)
        self._render_content_blocks(doc, content)

    def _render_content_blocks(self, doc: Document, content: str) -> None:
        for block in (content or "").split("\n\n"):
            block = normalize_unicode_text(block.strip())
            if not block:
                continue
            if block.startswith("## "):
                sub = doc.add_heading(block[3:].strip(), level=2)
                sub.runs[0].font.color.rgb = RGBColor(70, 130, 180)
            elif is_markdown_table(block):
                self._add_markdown_table(doc, block)
            elif is_bullet_block(block):
                for line in block.splitlines():
                    item = line[2:].strip()
                    doc.add_paragraph(item, style="List Bullet")
            elif block.startswith("[") and "Placeholder" in block:
                self._add_placeholder_box(doc, block)
            else:
                p = doc.add_paragraph(block)
                p.paragraph_format.space_after = Pt(6)

    def _add_markdown_table(self, doc: Document, markdown_block: str) -> None:
        rows = parse_markdown_table(markdown_block)
        if not rows:
            return
        table = doc.add_table(rows=1, cols=len(rows[0]))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        header_cells = table.rows[0].cells
        for i, cell in enumerate(rows[0]):
            header_cells[i].text = normalize_unicode_text(cell)
            self._shade_cell(header_cells[i], "2C5AA0")
            for paragraph in header_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)

        for row in rows[1:]:
            cells = table.add_row().cells
            for i, cell in enumerate(row):
                cells[i].text = normalize_unicode_text(cell)

    def _add_placeholder_box(self, doc: Document, text: str) -> None:
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.rows[0].cells[0]
        cell.text = normalize_unicode_text(text)
        self._shade_cell(cell, "FFF8E7")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)
        table.autofit = True

    def _add_placeholders(self, doc: Document, placeholders: Dict) -> None:
        heading = doc.add_heading("Draft Placeholders", level=1)
        heading.runs[0].font.color.rgb = RGBColor(44, 90, 160)
        for item in placeholder_items(placeholders):
            box = doc.add_table(rows=1, cols=1)
            box.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = box.rows[0].cells[0]
            cell.text = f"{item['label']}\n{item['instruction']}"
            self._shade_cell(cell, "FFF8E7")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)

    def _add_references(self, doc: Document, references: str) -> None:
        heading = doc.add_heading("References", level=1)
        heading.runs[0].font.color.rgb = RGBColor(44, 90, 160)
        for ref_line in references.split("\n"):
            ref_line = ref_line.strip()
            if ref_line:
                p = doc.add_paragraph(ref_line)
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.hanging_indent = Inches(-0.25)
                p.paragraph_format.space_after = Pt(4)

    def _add_disclaimer(self, doc: Document) -> None:
        heading = doc.add_heading("Regulatory Disclaimer", level=1)
        heading.runs[0].font.color.rgb = RGBColor(44, 90, 160)
        disclaimer_text = (
            "This product monograph was automatically generated using artificial intelligence and information from "
            "public medical databases. It is a draft summary only and may not reflect the most current regulatory "
            "status. All claims should be independently verified against original sources and regulatory authorities. "
            "This document must be reviewed by a qualified medical professional before distribution."
        )
        p = doc.add_paragraph(disclaimer_text)
        p.paragraph_format.space_after = Pt(6)
        footer = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %H:%M UTC')}")
        footer.runs[0].font.italic = True
        footer.runs[0].font.size = Pt(9)

    def _shade_cell(self, cell, color: str) -> None:
        shading_elm = OxmlElement("w:shd")
        shading_elm.set(qn("w:fill"), color)
        cell._element.get_or_add_tcPr().append(shading_elm)


class GoogleDocsGenerator:
    """Generates a text import template for Google Docs."""

    def __init__(self, output_dir: str = "data/monographs"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def create_google_docs_template(self, monograph_data: Dict) -> str:
        filename = f"{monograph_data.get('molecule_name', 'Monograph')}_GoogleDocs_Import_{datetime.now().strftime('%Y%m%d')}.txt"
        filepath = os.path.join(self.output_dir, filename)

        placeholders = []
        for item in placeholder_items(monograph_data.get("draft_placeholders")):
            placeholders.append(f"- {item['label']}: {item['instruction']}")

        content = f"""# IMPORT INSTRUCTIONS FOR GOOGLE DOCS

Molecule: {monograph_data.get('molecule_name')}
Generated: {datetime.now().strftime('%B %d, %Y')}

## STEPS TO CREATE IN GOOGLE DOCS:

1. Create a new Google Doc.
2. Paste the content below.
3. Apply heading styles and tables as needed.
4. Review with medical and regulatory stakeholders.

## DOCUMENT CONTENT:

=== START OF CONTENT ===

# {monograph_data.get('molecule_name')} - PRODUCT MONOGRAPH

## EXECUTIVE SUMMARY
{monograph_data.get('executive_summary', 'N/A')}

## MAIN SECTIONS

{self._format_sections(monograph_data)}

## DRAFT PLACEHOLDERS
{chr(10).join(placeholders) if placeholders else 'N/A'}

## REFERENCES
{monograph_data.get('references', 'N/A')}

=== END OF CONTENT ===
"""

        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)

        print(f"[OK] Google Docs import template: {filepath}")
        return filepath

    def _format_sections(self, monograph_data: Dict) -> str:
        formatted = ""
        for section_name, section_content in monograph_data.get("sections", {}).items():
            formatted += f"\n### {section_name.replace('_', ' ').title()}\n\n{normalize_unicode_text(section_content)}\n\n"
        return formatted


word_generator = WordDocumentGenerator()
google_docs_generator = GoogleDocsGenerator()
